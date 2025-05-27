
import os
import json
import numpy as np
import pickle
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

import faiss
from sentence_transformers import SentenceTransformer

import unicodedata
import re

import PyPDF2
from docx import Document
import openpyxl
import csv
import xml.etree.ElementTree as ET


class AdvancedDocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.xlsx', '.txt', '.json', '.csv', '.xml']
        self.logger = logging.getLogger(__name__)

    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""

    def _remove_extract_text_from_pptx(self, file_path: str) -> str:
        try:
            prs = pptx.Presentation(file_path)
            return "\n".join(shape.text for slide in prs.slides for shape in slide.shapes if hasattr(shape, "text"))
        except Exception as e:
            self.logger.error(f"Error reading PPTX {file_path}: {e}")
            return ""

    def extract_text_from_xlsx(self, file_path: str) -> str:
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            content = ""
            for sheet in workbook.sheetnames:
                ws = workbook[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                    content += row_text + "\n"
            return content
        except Exception as e:
            self.logger.error(f"Error reading XLSX {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            self.logger.error(f"Error reading TXT {file_path}: {e}")
            return ""

    def extract_text_from_json(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return self._json_to_text(data)
        except Exception as e:
            self.logger.error(f"Error reading JSON {file_path}: {e}")
            return ""

    def extract_text_from_csv(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return "\n".join(" | ".join(row) for row in csv.reader(f))
        except Exception as e:
            self.logger.error(f"Error reading CSV {file_path}: {e}")
            return ""

    def extract_text_from_xml(self, file_path: str) -> str:
        try:
            root = ET.parse(file_path).getroot()
            return self._xml_to_text(root)
        except Exception as e:
            self.logger.error(f"Error reading XML {file_path}: {e}")
            return ""

    def _json_to_text(self, data) -> str:
        if isinstance(data, dict):
            return "\n".join(f"{k}: {self._json_to_text(v)}" for k, v in data.items())
        elif isinstance(data, list):
            return "\n".join(self._json_to_text(i) for i in data)
        else:
            return str(data)

    def _xml_to_text(self, element) -> str:
        text = element.text or ""
        for child in element:
            text += self._xml_to_text(child)
        return text

    def process_file(self, file_path: str) -> Dict[str, Any]:
        ext = Path(file_path).suffix.lower()
        if ext not in self.supported_extensions:
            self.logger.warning(f"Unsupported type: {ext}")
            return None

        extractors = {
            ".pdf": self.extract_text_from_pdf,
            ".docx": self.extract_text_from_docx,
            
            ".xlsx": self.extract_text_from_xlsx,
            ".txt": self.extract_text_from_txt,
            ".json": self.extract_text_from_json,
            ".csv": self.extract_text_from_csv,
            ".xml": self.extract_text_from_xml
        }

        content = extractors[ext](file_path)
        if not content:
            return None

        stat = os.stat(file_path)
        return {
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "file_size": stat.st_size,
            "file_type": ext,
            "created_time": datetime.fromtimestamp(stat.st_ctime),
            "modified_time": datetime.fromtimestamp(stat.st_mtime),
            "content": content,
            "content_length": len(content)
        }


class VietnameseTextProcessor:
    def normalize_text(self, text: str) -> str:
        text = unicodedata.normalize("NFC", text)
        text = re.sub(r"[^\w\s]", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def tokenize(self, text: str) -> List[str]:
        try:
            return text.split()
        except:
            return text.split()

    def chunk_text(self, text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
        words = self.tokenize(text)
        return [
            " ".join(words[i:i + chunk_size])
            for i in range(0, len(words), chunk_size - overlap)
        ]


class AdvancedVectorStore:
    def __init__(self, model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"):
        self.model = SentenceTransformer(model_name)
        self.index = faiss.IndexFlatIP(self.model.get_sentence_embedding_dimension())
        self.documents = []
        self.chunks = []
        self.model_name = model_name

    
        BANNED_VECTOR_KEYWORDS = ["vương đình huệ", "nguyễn phú trọng", "chủ tịch quốc hội", "tổng bí thư", "quốc hội"]
    def add_documents(self, documents: List[Dict[str, Any]]):
        processor = VietnameseTextProcessor()
        for doc in documents:
            normalized = processor.normalize_text(doc["content"])
            chunks = processor.chunk_text(normalized)

            for i, chunk in enumerate(chunks):
                if any(banned in chunk.lower() for banned in BANNED_VECTOR_KEYWORDS):
                    continue  # Bỏ đoạn nhạy cảm

                emb = self.model.encode([chunk])[0]
                emb /= np.linalg.norm(emb)
                self.index.add(np.array([emb], dtype=np.float32))
                self.chunks.append({
                    "chunk_id": len(self.chunks),
                    "chunk_index": i,
                    "chunk_text": chunk,
                    "document_id": len(self.documents),
                    "file_path": doc.get("file_path"),
                    "file_name": doc.get("file_name"),
                })
            doc["chunk_count"] = len(chunks)
            self.documents.append(doc)

    def search(self, query: str, k=5, score_threshold=0.5) -> List[Dict[str, Any]]:
        if self.index.ntotal == 0:
            return []
        processor = VietnameseTextProcessor()
        query = processor.normalize_text(query)
        emb = self.model.encode([query])[0]
        emb /= np.linalg.norm(emb)
        scores, indices = self.index.search(np.array([emb], dtype=np.float32), min(k, self.index.ntotal))
        return [
            {**self.chunks[i], "similarity_score": float(scores[0][idx])}
            for idx, i in enumerate(indices[0]) if scores[0][idx] >= score_threshold
        ]

    def save(self, directory: str):
        os.makedirs(directory, exist_ok=True)
        faiss.write_index(self.index, os.path.join(directory, "vector_index.faiss"))
        with open(os.path.join(directory, "documents.pkl"), "wb") as f:
            pickle.dump(self.documents, f)
        with open(os.path.join(directory, "chunks.pkl"), "wb") as f:
            pickle.dump(self.chunks, f)

    def load(self, directory: str):
        self.index = faiss.read_index(os.path.join(directory, "vector_index.faiss"))
        with open(os.path.join(directory, "documents.pkl"), "rb") as f:
            self.documents = pickle.load(f)
        with open(os.path.join(directory, "chunks.pkl"), "rb") as f:
            self.chunks = pickle.load(f)


class DocumentRAGSystem:
    def __init__(self, vector_store_path: str = None):
        self.document_processor = AdvancedDocumentProcessor()
        self.vector_store = AdvancedVectorStore()
        if vector_store_path and os.path.exists(vector_store_path):
            self.vector_store.load(vector_store_path)

    def process_directory(self, directory: str):
        docs = []
        for file_path in Path(directory).rglob("*"):
            if file_path.is_file():
                doc_data = self.document_processor.process_file(str(file_path))
                if doc_data:
                    docs.append(doc_data)
        if docs:
            self.vector_store.add_documents(docs)
        return docs
