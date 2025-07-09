# server/services/vector_rag/document_processor.py
"""
Document Processor for Legal RAG Chatbot
Xử lý 2 loại văn bản:
1. Luật/Thông tư/Nghị định (.docx) - tách theo Điều/Khoản/Điểm
2. Văn bản Hỏi-Đáp (.docx) - tách theo cặp Hỏi:Đáp:
"""
import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from services.vector_rag.rag_config import config

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    

@dataclass
class Document:
    """Document data structure"""
    content: str
    metadata: Dict[str, Any]

class DocumentProcessor:
    """Document Processor cho Legal RAG Chatbot"""
    
    def __init__(self):
        self.documents_path = config.documents_path
        # Legal document patterns
        self.legal_patterns = {
            'law_number': re.compile(r'(?:Luật\s+số\s+)?(\d+/\d{4}/QH\d+)', re.IGNORECASE),
            'decree_number': re.compile(r'(?:Nghị\s+định\s+số\s+)?(\d+/\d{4}/NĐ-CP)', re.IGNORECASE),
            'circular_number': re.compile(r'(?:Thông\s+tư\s+số\s+)?(\d+/\d{4}/TT-[\w]+)', re.IGNORECASE),
            'doc_title': re.compile(r'(?:LUẬT|NGHỊ ĐỊNH|THÔNG TƯ)\s+(.+?)(?:\n|$)', re.IGNORECASE),
            'article': re.compile(r'Điều\s+(\d+[a-z]?)\s*[.:]?\s*(.*?)(?=\nĐiều|\nChương|\nMục|\Z)', re.DOTALL | re.IGNORECASE),
            'paragraph': re.compile(r'^\s*(\d+)\.\s+(.*?)(?=\n\d+\.|\nĐiều|\Z)', re.DOTALL | re.MULTILINE),
            'point': re.compile(r'^\s*([a-z])\)\s+(.*?)(?=\n[a-z]\)|\n\d+\.|\nĐiều|\Z)', re.DOTALL | re.MULTILINE)
        }
        
    
        self.qa_patterns = {
            'qa_pair': re.compile(r'(?:>\s*)?\*\*Hỏi:\s*(.*?)\*\*\s*(?:>\s*)?\*\*Đáp:\*\*(.*?)(?=(?:>\s*)?\*\*Hỏi:|\Z)', re.DOTALL | re.IGNORECASE),
            'simple_qa': re.compile(r'(?:>\s*)?Hỏi:\s*(.*?)(?:>\s*)?Đáp:\s*(.*?)(?=(?:>\s*)?Hỏi:|\Z)', re.DOTALL | re.IGNORECASE)
        }
    
    def process_file(self, file_path: str) -> List[Document]:
        """Xử lý một file .docx"""
        if not DOCX_AVAILABLE:
            return []
        
        if not file_path.endswith('.docx'):
            return []
        
        if not os.path.exists(file_path):
            return []
        
        try:
            # Extract content từ DOCX
            content = self._extract_docx_content(file_path)
            if not content or len(content) < 100:
                print(f"⚠️ {file_path} is too short or empty.")
                return []
            
            # Phân loại document
            doc_type = self._classify_document(content)
            
            if doc_type == 'legal':
                return self._process_legal_document(content, file_path)
            elif doc_type == 'qa':
                return self._process_qa_document(content, file_path)
            else:
                print(f"⚠️ {file_path} could not be classified (not legal or qa document).")
                return []  # Bỏ qua nếu không nhận diện được
                
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return []
    
    def process_directory(self, documents_path: str) -> List[Document]:
        if documents_path is None:
            documents_path = self.documents_path
        """Xử lý tất cả file .docx trong thư mục"""
        all_documents = []
        
        if not os.path.exists(documents_path):
            return []
        
        # Tìm tất cả file .docx
        docx_files = list(Path(documents_path).rglob('*.docx'))
        
        for file_path in docx_files:
            documents = self.process_file(str(file_path))
            all_documents.extend(documents)
        
        return all_documents
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extract nội dung từ file DOCX"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            return '\n'.join(paragraphs)
        except:
            return ""
    
    def _classify_document(self, content: str) -> str:
        """Phân loại document: legal, qa, hoặc unknown"""
        # Đếm số Điều
        dieu_count = len(re.findall(r'Điều\s+\d+', content, re.IGNORECASE))
        
        # Đếm số cặp Hỏi-Đáp
        qa_count = 0
        qa_count += len(re.findall(r'\*\*Hỏi:\*\*.*?\*\*Đáp:\*\*', content, re.DOTALL | re.IGNORECASE))
        qa_count += len(re.findall(r'Hỏi:.*?Đáp:', content, re.DOTALL | re.IGNORECASE))
        
        # DEBUG (optional)
        print(f"DEBUG: Điều count: {dieu_count}, Q&A count: {qa_count}")
        
        # ƯU TIÊN Q&A TRƯỚC (FIX)
        if qa_count >= 2:
            print("DEBUG: Classified as QA")
            return 'qa'
        elif dieu_count >= 3:
            print("DEBUG: Classified as LEGAL")
            return 'legal'
        else:
            print("DEBUG: Classified as UNKNOWN")
            return 'unknown'
    
    def _process_legal_document(self, content: str, file_path: str) -> List[Document]:
        """Xử lý văn bản pháp luật - tách theo Điều/Khoản/Điểm"""
        documents = []
        file_name = os.path.basename(file_path)
        
        # Extract document metadata
        doc_info = self._extract_legal_metadata(content, file_name)
        
        # Tìm tất cả các Điều
        articles = list(self.legal_patterns['article'].finditer(content))
        
        chunk_id = 0
        
        for article_match in articles:
            article_number = article_match.group(1)
            article_content = article_match.group(2).strip()
            
            # Tách title và content của điều
            lines = article_content.split('\n')
            article_title = lines[0].strip() if lines else ""
            article_body = '\n'.join(lines[1:]).strip() if len(lines) > 1 else ""
            
            # Tìm các khoản trong điều này
            paragraphs = list(self.legal_patterns['paragraph'].finditer(article_body))
            
            if paragraphs:
                # Có khoản → tách từng khoản
                for para_match in paragraphs:
                    para_number = para_match.group(1)
                    para_content = para_match.group(2).strip()
                    
                    # Tìm các điểm trong khoản này
                    points = list(self.legal_patterns['point'].finditer(para_content))
                    
                    if points and len(para_content) > 400:  # Khoản dài và có nhiều điểm → tách điểm
                        for point_match in points:
                            point_letter = point_match.group(1)
                            point_content = point_match.group(2).strip()
                            
                            # Tạo chunk cho điểm
                            chunk_content = self._format_legal_chunk(
                                doc_info, article_number, article_title, 
                                para_number, point_letter, point_content
                            )
                            
                            chunk_metadata = self._create_legal_metadata(
                                doc_info, file_name, f"{article_number}.{para_number}.{point_letter}",
                                chunk_id
                            )
                            
                            documents.append(Document(chunk_content, chunk_metadata))
                            chunk_id += 1
                    else:
                        # Tạo chunk cho cả khoản
                        chunk_content = self._format_legal_chunk(
                            doc_info, article_number, article_title, 
                            para_number, None, para_content
                        )
                        
                        chunk_metadata = self._create_legal_metadata(
                            doc_info, file_name, f"{article_number}.{para_number}",
                            chunk_id
                        )
                        
                        documents.append(Document(chunk_content, chunk_metadata))
                        chunk_id += 1
            else:
                # Không có khoản → tạo chunk cho cả điều
                full_article = f"{article_title}\n{article_body}".strip()
                
                chunk_content = self._format_legal_chunk(
                    doc_info, article_number, article_title, 
                    None, None, full_article
                )
                
                chunk_metadata = self._create_legal_metadata(
                    doc_info, file_name, article_number,
                    chunk_id
                )
                
                documents.append(Document(chunk_content, chunk_metadata))
                chunk_id += 1
        
        # Update total chunks cho tất cả metadata
        for doc in documents:
            doc.metadata['total_chunks'] = len(documents)
        
        return documents
    
    def _process_qa_document(self, content: str, file_path: str) -> List[Document]:
        """Xử lý văn bản Hỏi-Đáp"""
        documents = []
        file_name = os.path.basename(file_path)
        
        # Thử pattern có ** trước
        qa_pairs = list(self.qa_patterns['qa_pair'].finditer(content))
        
        # Nếu không có, thử pattern đơn giản
        if not qa_pairs:
            qa_pairs = list(self.qa_patterns['simple_qa'].finditer(content))
        
        for i, qa_match in enumerate(qa_pairs):
            question = qa_match.group(1).strip()
            answer = qa_match.group(2).strip()
            
            # Bỏ qua nếu quá ngắn
            if len(question) < 10 or len(answer) < 20:
                continue
            
            # Format Q&A content
            qa_content = f"CÂU HỎI: {question}\n\nTRẢ LỜI: {answer}"
            
            # Tạo metadata
            qa_metadata = {
                'content_type': 'qa_entry',
                'source': file_name,
                'qa_index': i + 1,
                'processed_at': datetime.now().isoformat()
            }
            
            documents.append(Document(qa_content, qa_metadata))
        
        return documents
    
    def _extract_legal_metadata(self, content: str, file_name: str) -> Dict[str, Any]:
        """Extract metadata từ văn bản pháp luật"""
        doc_info = {
            'doc_id': '',
            'doc_type': 'legal_document',
            'name': '',
            'doc_category': 'law'  # default
        }
        
        # Tìm số văn bản
        law_match = self.legal_patterns['law_number'].search(content)
        decree_match = self.legal_patterns['decree_number'].search(content)
        circular_match = self.legal_patterns['circular_number'].search(content)
        
        if law_match:
            doc_info['doc_id'] = law_match.group(1)
            doc_info['doc_category'] = 'law'
        elif decree_match:
            doc_info['doc_id'] = decree_match.group(1)
            doc_info['doc_category'] = 'decree'
        elif circular_match:
            doc_info['doc_id'] = circular_match.group(1)
            doc_info['doc_category'] = 'circular'
        else:
            # Fallback to filename
            doc_info['doc_id'] = file_name.replace('.docx', '')
        
        # Tìm tên văn bản
        title_match = self.legal_patterns['doc_title'].search(content)
        if title_match:
            doc_info['name'] = title_match.group(1).strip()
        else:
            doc_info['name'] = file_name.replace('.docx', '').replace('_', ' ')
        
        return doc_info
    
    def _format_legal_chunk(self, doc_info: Dict, article_num: str, article_title: str,
                           para_num: Optional[str], point_letter: Optional[str], 
                           content: str) -> str:
        """Format chunk content với header"""
        # Tạo header
        header_parts = [f"[{doc_info['doc_id']} - {doc_info['name']}]"]
        header_parts.append(f"Điều {article_num}. {article_title}")
        
        if para_num:
            header_parts.append(f"Khoản {para_num}")
        
        if point_letter:
            header_parts.append(f"Điểm {point_letter}")
        
        header = " ".join(header_parts)
        
        return f"{header}\n\n{content}"
    
    def _create_legal_metadata(self, doc_info: Dict, file_name: str, law_unit: str,
                              chunk_id: int) -> Dict[str, Any]:
        """Tạo metadata cho legal chunk"""
        return {
            'content_type': 'legal_document',
            'doc_type': doc_info['doc_category'],
            'doc_id': doc_info['doc_id'],
            'name': doc_info['name'],
            'law_unit': law_unit,
            'source': file_name,
            'chunk_id': chunk_id,
            'total_chunks': 0,  # Sẽ được update sau
            'processed_at': datetime.now().isoformat()
        }
    
    def get_stats(self) -> Dict[str, Any]:
        """Lấy thống kê xử lý"""
        return {
            'processor_version': '1.0',
            'supported_formats': ['.docx'],
            'document_types': ['legal', 'qa'],
            'features': [
                'Legal structure extraction (Điều/Khoản/Điểm)',
                'Q&A pair extraction',
                'Smart chunking with metadata',
                'Document classification'
            ]
        }

# Factory function
def create_processor() -> DocumentProcessor:
    """Tạo DocumentProcessor instance"""
    return DocumentProcessor()