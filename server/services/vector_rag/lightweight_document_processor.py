# server/services/vector_rag/lightweight_document_processor.py
"""
Document processor siêu nhẹ, tối ưu cho văn bản pháp lý xuất nhập cảnh
Thay thế LangChain document loaders bằng implementation nhẹ hơn
"""
import os
import re
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import logging

# Optional imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Thay đổi import ở đầu file:
from .lightweight_config import CHUNKING_CONFIG, SYSTEM_CONFIG, DOCUMENT_TYPES

logger = logging.getLogger(__name__)

class LegalDocument:
    """Lightweight document class thay thế LangChain Document"""
    
    def __init__(self, content: str, metadata: Dict[str, Any] = None):
        self.page_content = content
        self.metadata = metadata or {}

class LightweightDocumentProcessor:
    """
    Document processor tối ưu cho văn bản pháp lý
    Không phụ thuộc LangChain, xử lý nhanh hơn
    """
    
    def __init__(self, config=None, system_config=None):
        self.config = config or CHUNKING_CONFIG
        self.system_config = system_config or SYSTEM_CONFIG
        
        # Cache cho file đã xử lý
        self.processed_cache = {}
        self.cache_file = os.path.join(
            self.system_config.cache_path, 
            "processed_docs_cache.json"
        )
        
        self._load_processed_cache()
    
    def _load_processed_cache(self):
        """Load cache của files đã xử lý"""
        try:
            import json
            if os.path.exists(self.cache_file):
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    self.processed_cache = json.load(f)
                logger.info(f"📂 Loaded cache for {len(self.processed_cache)} processed files")
        except Exception as e:
            logger.warning(f"⚠️ Failed to load processed cache: {e}")
            self.processed_cache = {}
    
    def _save_processed_cache(self):
        """Lưu cache files đã xử lý"""
        try:
            import json
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.processed_cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"⚠️ Failed to save processed cache: {e}")
    
    def _get_file_hash(self, file_path: str) -> str:
        """Tính hash của file để detect thay đổi"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            return hashlib.md5(content).hexdigest()
        except Exception:
            return ""
    
    def _is_file_cached(self, file_path: str) -> bool:
        """Kiểm tra file đã được xử lý chưa"""
        file_hash = self._get_file_hash(file_path)
        cached_hash = self.processed_cache.get(file_path, {}).get('hash', '')
        return file_hash == cached_hash and file_hash != ""
    
    def _update_file_cache(self, file_path: str, doc_count: int):
        """Cập nhật cache cho file"""
        file_hash = self._get_file_hash(file_path)
        self.processed_cache[file_path] = {
            'hash': file_hash,
            'processed_at': datetime.now().isoformat(),
            'doc_count': doc_count
        }
    
    def load_txt_file(self, file_path: str) -> str:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback encoding
            with open(file_path, 'r', encoding='cp1252', errors='ignore') as f:
                return f.read()
    
    def load_pdf_file(self, file_path: str) -> str:
        """Load PDF file"""
        if not PDF_AVAILABLE:
            logger.warning(f"⚠️ PyPDF2 not available, skipping {file_path}")
            return ""
        
        try:
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    try:
                        text += page.extract_text() + "\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract page: {e}")
                        continue
            return text
        except Exception as e:
            logger.error(f"❌ Failed to load PDF {file_path}: {e}")
            return ""
    
    def load_docx_file(self, file_path: str) -> str:
        """Load DOCX file"""
        if not DOCX_AVAILABLE:
            logger.warning(f"⚠️ python-docx not available, skipping {file_path}")
            return ""
        
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"❌ Failed to load DOCX {file_path}: {e}")
            return ""
    
    def detect_document_type(self, file_path: str, content: str) -> str:
        """Detect loại văn bản pháp lý từ tên file và nội dung"""
        file_name = Path(file_path).name.lower()
        content_lower = content.lower()
        
        # Check tên file trước
        for doc_type, type_name in DOCUMENT_TYPES.items():
            if doc_type in file_name:
                return doc_type
        
        # Check nội dung
        if 'luật' in content_lower and 'số' in content_lower:
            return 'luat'
        elif 'nghị định' in content_lower:
            return 'nghidinh'
        elif 'thông tư' in content_lower:
            return 'thongtu'
        elif 'quyết định' in content_lower:
            return 'quyetdinh'
        elif 'hướng dẫn' in content_lower:
            return 'huongdan'
        
        return 'general'
    
    def extract_legal_references(self, content: str) -> List[str]:
        """Trích xuất các tham chiếu pháp lý (Điều, Khoản, Điểm)"""
        references = []
        
        # Pattern cho Điều luật
        dieu_pattern = r'Điều\s+\d+[a-z]?\.?'
        references.extend(re.findall(dieu_pattern, content, re.IGNORECASE))
        
        # Pattern cho Khoản
        khoan_pattern = r'Khoản\s+\d+[a-z]?\.?'
        references.extend(re.findall(khoan_pattern, content, re.IGNORECASE))
        
        # Pattern cho Điểm
        diem_pattern = r'Điểm\s+[a-z]+\)'
        references.extend(re.findall(diem_pattern, content, re.IGNORECASE))
        
        return list(set(references))  # Remove duplicates
    
    def legal_aware_chunking(self, content: str) -> List[str]:
        """
        Chunking thông minh cho văn bản pháp lý
        Ưu tiên giữ nguyên cấu trúc Điều, Khoản, Điểm
        """
        chunks = []
        
        # Split theo Điều trước
        dieu_splits = re.split(r'(Điều\s+\d+[a-z]?\.?\s*)', content, flags=re.IGNORECASE)
        
        current_chunk = ""
        for i, part in enumerate(dieu_splits):
            if re.match(r'Điều\s+\d+', part, re.IGNORECASE):
                # Bắt đầu Điều mới
                if current_chunk and len(current_chunk) > 100:
                    chunks.extend(self._split_chunk_if_needed(current_chunk))
                current_chunk = part
            else:
                current_chunk += part
                
                # Check size
                if len(current_chunk) > self.config.chunk_size:
                    chunks.extend(self._split_chunk_if_needed(current_chunk))
                    current_chunk = ""
        
        # Add chunk cuối
        if current_chunk.strip():
            chunks.extend(self._split_chunk_if_needed(current_chunk))
        
        return [chunk.strip() for chunk in chunks if chunk.strip()]
    
    def _split_chunk_if_needed(self, chunk: str) -> List[str]:
        """Split chunk nếu quá dài"""
        if len(chunk) <= self.config.chunk_size:
            return [chunk]
        
        # Split theo separators
        for separator in self.config.separators:
            if separator in chunk:
                parts = chunk.split(separator)
                result = []
                current = ""
                
                for part in parts:
                    if len(current + separator + part) <= self.config.chunk_size:
                        current += separator + part if current else part
                    else:
                        if current:
                            result.append(current)
                        current = part
                
                if current:
                    result.append(current)
                
                return result
        
        # Force split nếu không có separator phù hợp
        return [chunk[i:i+self.config.chunk_size] 
                for i in range(0, len(chunk), self.config.chunk_size)]
    
    def process_single_file(self, file_path: str) -> List[LegalDocument]:
        """Xử lý một file duy nhất"""
        file_ext = Path(file_path).suffix.lower()
        
        # Check cache
        if self._is_file_cached(file_path):
            logger.info(f"📋 File cached, skipping: {Path(file_path).name}")
            return []
        
        logger.info(f"📄 Processing: {Path(file_path).name}")
        
        # Load content dựa theo extension
        content = ""
        if file_ext == '.txt':
            content = self.load_txt_file(file_path)
        elif file_ext == '.pdf':
            content = self.load_pdf_file(file_path)
        elif file_ext == '.docx':
            content = self.load_docx_file(file_path)
        else:
            logger.warning(f"⚠️ Unsupported file type: {file_ext}")
            return []
        
        if not content.strip():
            logger.warning(f"⚠️ Empty content: {file_path}")
            return []
        
        # Detect document type
        doc_type = self.detect_document_type(file_path, content)
        
        # Extract legal references
        legal_refs = self.extract_legal_references(content)
        
        # Chunking
        chunks = self.legal_aware_chunking(content)
        
        # Tạo documents
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                'source': file_path,
                'file_type': file_ext,
                'doc_type': doc_type,
                'chunk_id': i,
                'chunk_size': len(chunk),
                'legal_references': legal_refs,
                'domain': 'xuatnhapcanh',
                'processed_at': datetime.now().isoformat()
            }
            
            documents.append(LegalDocument(content=chunk, metadata=metadata))
        
        # Update cache
        self._update_file_cache(file_path, len(documents))
        
        logger.info(f"✅ Processed {len(documents)} chunks from {Path(file_path).name}")
        return documents
    
    def process_documents_directory(self, directory_path: str = None) -> List[LegalDocument]:
        """Xử lý tất cả documents trong thư mục"""
        if directory_path is None:
            directory_path = self.system_config.documents_path
        
        if not os.path.exists(directory_path):
            logger.error(f"❌ Directory not found: {directory_path}")
            return []
        
        logger.info(f"📁 Processing documents from: {directory_path}")
        
        all_documents = []
        processed_files = 0
        skipped_files = 0
        
        # Scan tất cả files trong directory
        for file_path in Path(directory_path).rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.system_config.supported_formats:
                try:
                    docs = self.process_single_file(str(file_path))
                    if docs:
                        all_documents.extend(docs)
                        processed_files += 1
                    else:
                        skipped_files += 1
                        
                except Exception as e:
                    logger.error(f"❌ Failed to process {file_path}: {e}")
                    skipped_files += 1
        
        # Save cache
        self._save_processed_cache()
        
        logger.info(f"✅ Processing complete:")
        logger.info(f"   📄 Processed files: {processed_files}")
        logger.info(f"   ⏭️ Skipped files: {skipped_files}")
        logger.info(f"   📝 Total chunks: {len(all_documents)}")
        
        return all_documents
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Thống kê quá trình xử lý"""
        stats = {
            'cached_files': len(self.processed_cache),
            'supported_formats': self.system_config.supported_formats,
            'chunk_size': self.config.chunk_size,
            'chunk_overlap': self.config.chunk_overlap,
            'cache_file': self.cache_file
        }
        
        # Thống kê theo doc type
        doc_types = {}
        for file_info in self.processed_cache.values():
            doc_count = file_info.get('doc_count', 0)
            # Estimate doc type from filename (simplified)
            for doc_type in DOCUMENT_TYPES.keys():
                if doc_type in str(file_info):
                    doc_types[doc_type] = doc_types.get(doc_type, 0) + doc_count
                    break
        
        stats['doc_types'] = doc_types
        return stats
    
    def clear_cache(self):
        """Xóa cache xử lý"""
        self.processed_cache = {}
        if os.path.exists(self.cache_file):
            os.remove(self.cache_file)
        logger.info("🗑️ Processing cache cleared")
    
    def reprocess_file(self, file_path: str) -> List[LegalDocument]:
        """Force reprocess một file cụ thể"""
        if file_path in self.processed_cache:
            del self.processed_cache[file_path]
        return self.process_single_file(file_path)

# Utility functions
def extract_xuatnhapcanh_entities(content: str) -> Dict[str, List[str]]:
    """Trích xuất entities liên quan xuất nhập cảnh"""
    entities = {
        'countries': [],
        'visa_types': [],
        'procedures': [],
        'timeframes': [],
        'fees': []
    }
    
    # Countries pattern
    country_pattern = r'(?:người|công dân)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
    entities['countries'] = re.findall(country_pattern, content)
    
    # Visa types
    visa_pattern = r'(?:visa|thị thực)\s+([a-záàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ\s]+)'
    entities['visa_types'] = re.findall(visa_pattern, content, re.IGNORECASE)
    
    # Timeframes
    time_pattern = r'(\d+\s*(?:ngày|tháng|năm))'
    entities['timeframes'] = re.findall(time_pattern, content)
    
    # Fees
    fee_pattern = r'(\d+(?:\.\d+)*\s*(?:đồng|USD|VND))'
    entities['fees'] = re.findall(fee_pattern, content)
    
    return entities

def validate_legal_document(content: str) -> Dict[str, Any]:
    """Validate tính hợp lệ của văn bản pháp lý"""
    validation = {
        'is_valid': True,
        'issues': [],
        'confidence': 1.0
    }
    
    # Check minimum length
    if len(content) < 100:
        validation['issues'].append("Content too short")
        validation['confidence'] -= 0.3
    
    # Check có structure pháp lý không
    if not re.search(r'Điều\s+\d+', content, re.IGNORECASE):
        validation['issues'].append("No legal articles found")
        validation['confidence'] -= 0.2
    
    # Check ngôn ngữ
    vietnamese_chars = 'áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ'
    if not any(char in content.lower() for char in vietnamese_chars):
        validation['issues'].append("Possible non-Vietnamese content")
        validation['confidence'] -= 0.1
    
    validation['is_valid'] = validation['confidence'] > 0.5
    return validation

# Test function
def test_document_processing():
    """Test document processing functionality"""
    processor = LightweightDocumentProcessor()
    
    # Test với sample text
    sample_content = """
    Điều 15. Thị thực nhập cảnh
    1. Người nước ngoài nhập cảnh vào Việt Nam phải có thị thực, trừ trường hợp được miễn thị thực theo quy định của pháp luật Việt Nam hoặc điều ước quốc tế mà Việt Nam là thành viên.
    2. Thị thực nhập cảnh được cấp cho các mục đích sau đây:
    a) Du lịch;
    b) Thăm thân;
    c) Công tác;
    d) Đầu tư, kinh doanh;
    đ) Lao động;
    e) Học tập;
    g) Khám bệnh, chữa bệnh;
    h) Tham gia hoạt động tư pháp;
    i) Mục đích khác.
    """
    
    print("🧪 Testing document processing...")
    
    # Test chunking
    chunks = processor.legal_aware_chunking(sample_content)
    print(f"✅ Generated {len(chunks)} chunks")
    
    # Test entity extraction
    entities = extract_xuatnhapcanh_entities(sample_content)
    print(f"📊 Extracted entities: {entities}")
    
    # Test validation
    validation = validate_legal_document(sample_content)
    print(f"✅ Validation: {validation}")
    
    return True

# if __name__ == "__main__":
#     test_document_processing()